from __future__ import annotations

import json
from datetime import UTC, datetime
from io import BytesIO
from uuid import uuid4

import pytest

from app.graph import (
    EntityMention,
    EntityRelationExtractor,
    EvidenceRelation,
    GraphBatch,
    GraphSearchRequest,
    KuzuGraphRepository,
)
from app.ingestion.chunking import DeterministicChunker
from app.ingestion.errors import PermanentIngestionError
from app.ingestion.events import DocumentIngestRequestedEvent
from app.ingestion.extraction import DocumentTextExtractor
from app.ingestion.spreadsheets import CalculationCommand, PolarsCalculationAdapter, TypedFilter
from app.rag.errors import ChatModelProviderError


def test_markdown_and_html_preserve_structure_and_ignore_active_content() -> None:
    extractor = DocumentTextExtractor()
    markdown = extractor.parse(
        b"# Policy\n\nIntro text.\n\n- First\n- Second\n\n```py\nprint('safe')\n```",
        content_type="text/markdown",
        file_name="policy.md",
    )
    html = extractor.parse(
        b"<h1>Policy</h1><script>steal()</script><p>Safe text</p><ul><li>One</li></ul>",
        content_type="text/html",
        file_name="policy.html",
    )

    assert [block.block_type for block in markdown.blocks] == [
        "heading",
        "paragraph",
        "list",
        "code",
    ]
    assert markdown.blocks[1].section_path == ("Policy",)
    assert "steal" not in " ".join(block.text for block in html.blocks)
    assert html.blocks[1].section_path == ("Policy",)


def test_csv_tables_types_ranges_and_calculations_are_deterministic() -> None:
    parsed = DocumentTextExtractor().parse(
        b"item;amount;day\nA;10.50;2026-07-01\nB;20.25;2026-07-02\n\ngroup;count\nX;2\n",
        content_type="text/csv",
        file_name="sales.csv",
    )

    assert parsed.modality == "spreadsheet"
    assert len(parsed.tables) == 2
    table = parsed.tables[0]
    assert [column.inferred_type for column in table.columns] == ["string", "decimal", "date"]
    assert table.rows[0].cell_range == "A2:C2"
    chunks = DeterministicChunker(chunk_size=80, overlap=10).chunk(parsed)
    assert all(chunk.table_id for chunk in chunks)
    assert all(chunk.cell_range for chunk in chunks)

    adapter = PolarsCalculationAdapter()
    result = adapter.execute(
        table,
        CalculationCommand(
            table_id=table.table_id,
            operation="sum",
            column="amount",
            filters=(TypedFilter(column="day", operator="gte", value="2026-07-01"),),
        ),
    )
    assert float(result.value) == pytest.approx(30.75)
    with pytest.raises(PermanentIngestionError, match="Unknown spreadsheet column"):
        adapter.execute(
            table,
            CalculationCommand(table_id=table.table_id, operation="sum", column="missing"),
        )


def test_docx_and_xlsx_extract_structures_and_exclude_formulas() -> None:
    from docx import Document
    from openpyxl import Workbook

    document = Document()
    document.add_heading("Returns", level=1)
    document.add_paragraph("Return within seven days.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 1).text = "Days"
    table.cell(1, 0).text = "Sale"
    table.cell(1, 1).text = "7"
    docx = BytesIO()
    document.save(docx)
    parsed_docx = DocumentTextExtractor().parse(
        docx.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        file_name="returns.docx",
    )
    assert {block.block_type for block in parsed_docx.blocks} >= {"heading", "paragraph", "table"}
    docx_chunks = DeterministicChunker(chunk_size=800, overlap=120).chunk(parsed_docx)
    table_chunk = next(chunk for chunk in docx_chunks if chunk.block_type == "table")
    assert table_chunk.text == "Item | Days\nSale | 7"

    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Sales"
    sheet.append(["item", "amount", "computed"])
    sheet.append(["A", 10, "=B2*2"])
    hidden = workbook.create_sheet("Hidden")
    hidden.sheet_state = "hidden"
    hidden.append(["secret", "value"])
    hidden.append(["x", 1])
    xlsx = BytesIO()
    workbook.save(xlsx)
    parsed_xlsx = DocumentTextExtractor().parse(
        xlsx.getvalue(),
        content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        file_name="sales.xlsx",
    )
    assert len(parsed_xlsx.tables) == 1
    row = parsed_xlsx.tables[0].rows[0]
    assert row.values["computed"] is None
    assert row.formula_expressions == {"computed": "=B2*2"}


def test_graph_requires_evidence_and_replacement_is_idempotent(tmp_path: object) -> None:
    repository = KuzuGraphRepository(str(tmp_path) + "/graph")
    batch = GraphBatch(
        tenant_id="tenant-a",
        knowledge_base_id="kb-a",
        source_id="source-a",
        source_name="policy.txt",
        units=({"unit_id": "unit-a", "text": "Acme has a policy", "page_number": 1},),
        entities=(
            EntityMention(
                name="Acme",
                normalized_name="acme",
                entity_type="organization",
                evidence_unit_id="unit-a",
            ),
        ),
    )
    repository.replace_source(batch)
    repository.replace_source(batch)
    assert (
        len(
            repository.search(
                GraphSearchRequest(tenant_id="tenant-a", knowledge_base_id="kb-a", query="Acme")
            )
        )
        == 1
    )
    assert (
        repository.search(
            GraphSearchRequest(tenant_id="tenant-b", knowledge_base_id="kb-a", query="Acme")
        )
        == []
    )

    with pytest.raises(ValueError, match="unknown evidence"):
        GraphBatch(
            tenant_id="tenant-a",
            knowledge_base_id="kb-a",
            source_id="bad",
            source_name="bad.txt",
            units=batch.units,
            entities=(
                EntityMention(
                    name="X", normalized_name="x", entity_type="thing", evidence_unit_id="missing"
                ),
            ),
            relations=(
                EvidenceRelation(
                    subject_normalized_name="x",
                    predicate="is",
                    object_normalized_name="x",
                    evidence_unit_id="missing",
                ),
            ),
        )


@pytest.mark.asyncio
async def test_graph_extraction_splits_batches_when_model_hits_output_limit() -> None:
    class LengthLimitedModel:
        def __init__(self) -> None:
            self.batch_sizes: list[int] = []

        async def complete(self, messages: list[dict[str, object]]) -> str:
            import json

            payload = json.loads(str(messages[1]["content"]))
            self.batch_sizes.append(len(payload))
            if len(payload) > 1:
                raise ChatModelProviderError(
                    "Model provider returned an empty response (finish_reason=length)"
                )
            return '{"entities":[],"relations":[]}'

    parsed = DocumentTextExtractor().parse(
        b"First paragraph.\n\nSecond paragraph.",
        content_type="text/plain",
        file_name="notes.txt",
    )
    chunks = DeterministicChunker().chunk(parsed)
    model = LengthLimitedModel()

    entities, relations = await EntityRelationExtractor(model, batch_size=4)._extract_batch(chunks)

    assert entities == []
    assert relations == []
    assert model.batch_sizes == [2, 1, 1]


@pytest.mark.asyncio
async def test_graph_extraction_retries_single_unit_after_output_limit() -> None:
    class InitiallyLengthLimitedModel:
        def __init__(self) -> None:
            self.calls = 0

        async def complete(self, messages: list[dict[str, object]]) -> str:
            del messages
            self.calls += 1
            if self.calls == 1:
                raise ChatModelProviderError(
                    "Model provider returned an empty response (finish_reason=length)"
                )
            return '{"entities":[],"relations":[]}'

    parsed = DocumentTextExtractor().parse(
        b"One paragraph.",
        content_type="text/plain",
        file_name="notes.txt",
    )
    chunks = DeterministicChunker().chunk(parsed)
    model = InitiallyLengthLimitedModel()

    entities, relations = await EntityRelationExtractor(model)._extract_batch(chunks)

    assert entities == []
    assert relations == []
    assert model.calls == 2


@pytest.mark.asyncio
async def test_graph_extraction_fails_after_single_unit_output_limit_retry() -> None:
    class AlwaysLengthLimitedModel:
        def __init__(self) -> None:
            self.calls = 0

        async def complete(self, messages: list[dict[str, object]]) -> str:
            del messages
            self.calls += 1
            raise ChatModelProviderError(
                "Model provider returned an empty response (finish_reason=length)"
            )

    parsed = DocumentTextExtractor().parse(
        b"One paragraph.",
        content_type="text/plain",
        file_name="notes.txt",
    )
    chunks = DeterministicChunker().chunk(parsed)
    model = AlwaysLengthLimitedModel()

    with pytest.raises(PermanentIngestionError, match="configured model output limit"):
        await EntityRelationExtractor(model)._extract_batch(chunks)

    assert model.calls == 2


@pytest.mark.asyncio
async def test_graph_extraction_discards_ungrounded_entities_and_relations() -> None:
    class UngroundedModel:
        async def complete(self, messages: list[dict[str, object]]) -> str:
            units = json.loads(str(messages[1]["content"]))
            unit_id = units[0]["unit_id"]
            return json.dumps(
                {
                    "entities": [
                        {
                            "name": "Acme",
                            "normalized_name": "acme",
                            "entity_type": "organization",
                            "evidence_unit_id": unit_id,
                        },
                        {
                            "name": "Policy",
                            "normalized_name": "policy",
                            "entity_type": "document",
                            "evidence_unit_id": unit_id,
                        },
                        {
                            "name": "Ghost",
                            "normalized_name": "ghost",
                            "entity_type": "unknown",
                            "evidence_unit_id": "invented-unit",
                        },
                    ],
                    "relations": [
                        {
                            "subject_normalized_name": "acme",
                            "predicate": "has",
                            "object_normalized_name": "policy",
                            "evidence_unit_id": unit_id,
                        },
                        {
                            "subject_normalized_name": "acme",
                            "predicate": "owns",
                            "object_normalized_name": "missing",
                            "evidence_unit_id": unit_id,
                        },
                        {
                            "subject_normalized_name": "ghost",
                            "predicate": "haunts",
                            "object_normalized_name": "acme",
                            "evidence_unit_id": unit_id,
                        },
                        {
                            "subject_normalized_name": "acme",
                            "predicate": "has",
                            "object_normalized_name": "policy",
                            "evidence_unit_id": "invented-unit",
                        },
                    ],
                }
            )

    parsed = DocumentTextExtractor().parse(
        b"Acme has a policy.",
        content_type="text/plain",
        file_name="notes.txt",
    )
    chunks = DeterministicChunker().chunk(parsed)
    event = DocumentIngestRequestedEvent(
        schema_version="1.0",
        event_id=uuid4(),
        job_id=uuid4(),
        tenant_id=uuid4(),
        knowledge_base_id=uuid4(),
        document_id=uuid4(),
        uploader_id=uuid4(),
        storage_key="documents/notes.txt",
        file_name="notes.txt",
        content_type="text/plain",
        file_size_bytes=18,
        occurred_at=datetime.now(UTC),
    )

    batch = await EntityRelationExtractor(UngroundedModel()).extract(event, chunks)

    assert [entity.normalized_name for entity in batch.entities] == ["acme", "policy"]
    assert len(batch.relations) == 1
    assert batch.relations[0].predicate == "has"
