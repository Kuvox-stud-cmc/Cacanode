from __future__ import annotations

import csv
import hashlib
import re
from abc import ABC, abstractmethod
from dataclasses import dataclass, field, replace
from datetime import date, datetime
from decimal import Decimal
from html.parser import HTMLParser
from io import BytesIO, StringIO
from pathlib import PurePosixPath
from typing import Any, Literal

from app.ingestion.errors import PermanentIngestionError

BlockType = Literal[
    "heading", "paragraph", "list", "code", "quote", "table", "page", "sheet", "row"
]


@dataclass(frozen=True, slots=True)
class ExtractedPage:
    """Compatibility view retained for callers that still consume page text."""

    page_number: int
    text: str


@dataclass(frozen=True, slots=True)
class KnowledgeBlock:
    unit_id: str
    block_type: BlockType
    text: str
    section_path: tuple[str, ...] = ()
    heading_context: str | None = None
    page_number: int | None = None
    sheet_name: str | None = None
    cell_range: str | None = None
    table_id: str | None = None
    source_start: int | None = None
    source_end: int | None = None
    content_hash: str = ""
    parser_version: str = "digital-v1"
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    modality: Literal["document", "spreadsheet"]
    blocks: tuple[KnowledgeBlock, ...]
    tables: tuple[NormalizedTable, ...] = ()


@dataclass(frozen=True, slots=True)
class NormalizedColumn:
    name: str
    inferred_type: Literal["integer", "decimal", "date", "datetime", "boolean", "string"]


@dataclass(frozen=True, slots=True)
class NormalizedRow:
    row_number: int
    cell_range: str
    values: dict[str, Any]
    formula_columns: tuple[str, ...] = ()
    formula_expressions: dict[str, str] = field(default_factory=dict)


@dataclass(frozen=True, slots=True)
class NormalizedTable:
    table_id: str
    sheet_name: str
    cell_range: str
    columns: tuple[NormalizedColumn, ...]
    rows: tuple[NormalizedRow, ...]


def _block(
    source: str,
    ordinal: int,
    block_type: BlockType,
    text: str,
    **kwargs: Any,
) -> KnowledgeBlock:
    normalized = text.strip()
    identity = (
        f"{source}:{ordinal}:{block_type}:{kwargs.get('page_number')}:{kwargs.get('sheet_name')}"
    )
    return KnowledgeBlock(
        unit_id=hashlib.sha256(identity.encode()).hexdigest()[:32],
        block_type=block_type,
        text=normalized,
        content_hash=hashlib.sha256(normalized.encode("utf-8")).hexdigest(),
        **kwargs,
    )


def _decode_utf8(data: bytes, label: str) -> str:
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise PermanentIngestionError(f"{label} document is not valid UTF-8") from exc
    if not text.strip():
        raise PermanentIngestionError("Document contains no extractable text")
    return text


class ParserAdapter(ABC):
    extensions: tuple[str, ...] = ()
    content_types: tuple[str, ...] = ()

    @abstractmethod
    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument: ...


class TextParser(ParserAdapter):
    extensions = (".txt",)
    content_types = ("text/plain",)

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        text = _decode_utf8(data, "TXT")
        blocks: list[KnowledgeBlock] = []
        offset = 0
        for ordinal, match in enumerate(re.finditer(r"\S(?:.*?\S)?(?=\n\s*\n|\Z)", text, re.S)):
            paragraph = match.group(0).strip()
            blocks.append(
                _block(
                    file_name,
                    ordinal,
                    "paragraph",
                    paragraph,
                    page_number=1,
                    source_start=match.start(),
                    source_end=match.end(),
                )
            )
            offset = match.end()
        if not blocks and text.strip():
            blocks.append(
                _block(
                    file_name,
                    0,
                    "paragraph",
                    text,
                    page_number=1,
                    source_start=0,
                    source_end=len(text),
                )
            )
        del offset
        return ParsedDocument("document", tuple(blocks))


class PdfParser(ParserAdapter):
    extensions = (".pdf",)
    content_types = ("application/pdf",)

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        if not data.startswith(b"%PDF-"):
            raise PermanentIngestionError("PDF document signature is invalid")
        try:
            from pypdf import PdfReader

            reader = PdfReader(BytesIO(data))
            if reader.is_encrypted:
                raise PermanentIngestionError("Encrypted PDF files are not supported")
        except PermanentIngestionError:
            raise
        except Exception as exc:
            raise PermanentIngestionError("PDF document could not be parsed") from exc
        blocks: list[KnowledgeBlock] = []
        for page_number, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            for paragraph in re.split(r"\n\s*\n", text):
                if paragraph.strip():
                    blocks.append(
                        _block(file_name, len(blocks), "page", paragraph, page_number=page_number)
                    )
        if not blocks:
            raise PermanentIngestionError(
                "PDF document contains no extractable text; scanned PDFs are not supported"
            )
        return ParsedDocument("document", tuple(blocks))


class DocxParser(ParserAdapter):
    extensions = (".docx",)
    content_types = ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",)

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        try:
            from docx import Document

            document = Document(BytesIO(data))
        except ImportError as exc:
            raise PermanentIngestionError("DOCX parser dependency is not installed") from exc
        except Exception as exc:
            raise PermanentIngestionError("DOCX document could not be parsed") from exc
        headings: list[str] = []
        blocks: list[KnowledgeBlock] = []
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            style = (paragraph.style.name if paragraph.style is not None else "").lower()
            if style.startswith("heading"):
                match = re.search(r"(\d+)", style)
                level = int(match.group(1)) if match else 1
                headings[level - 1 :] = [text]
                block_type: BlockType = "heading"
            elif "list" in style:
                block_type = "list"
            else:
                block_type = "paragraph"
            blocks.append(
                _block(
                    file_name,
                    len(blocks),
                    block_type,
                    text,
                    section_path=tuple(headings),
                    heading_context=headings[-1] if headings else None,
                )
            )
        for table_index, table in enumerate(document.tables, start=1):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            text = "\n".join(" | ".join(row) for row in rows if any(row))
            if text:
                blocks.append(
                    _block(
                        file_name,
                        len(blocks),
                        "table",
                        text,
                        section_path=tuple(headings),
                        table_id=f"table-{table_index}",
                    )
                )
        if not blocks:
            raise PermanentIngestionError("Document contains no extractable text")
        return ParsedDocument("document", tuple(blocks))


class MarkdownParser(ParserAdapter):
    extensions = (".md", ".markdown")
    content_types = ("text/markdown", "text/x-markdown")

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        text = _decode_utf8(data, "Markdown")
        blocks: list[KnowledgeBlock] = []
        headings: list[str] = []
        current: list[str] = []
        current_type: BlockType = "paragraph"
        in_code = False

        def flush() -> None:
            nonlocal current
            value = "\n".join(current).strip()
            if value:
                blocks.append(
                    _block(
                        file_name,
                        len(blocks),
                        current_type,
                        value,
                        section_path=tuple(headings),
                        heading_context=headings[-1] if headings else None,
                    )
                )
            current = []

        for line in text.splitlines():
            if line.lstrip().startswith("```"):
                flush()
                in_code = not in_code
                current_type = "code" if in_code else "paragraph"
                continue
            if in_code:
                current.append(line)
                continue
            heading = re.match(r"^(#{1,6})\s+(.+)$", line)
            if heading:
                flush()
                level, title = len(heading.group(1)), heading.group(2).strip()
                headings[level - 1 :] = [title]
                blocks.append(
                    _block(
                        file_name,
                        len(blocks),
                        "heading",
                        title,
                        section_path=tuple(headings),
                        heading_context=title,
                    )
                )
                current_type = "paragraph"
            elif re.match(r"^\s*(?:[-*+] |\d+[.)] )", line):
                if current_type != "list":
                    flush()
                    current_type = "list"
                current.append(line.strip())
            elif line.lstrip().startswith(">"):
                if current_type != "quote":
                    flush()
                    current_type = "quote"
                current.append(line.lstrip()[1:].strip())
            elif "|" in line and line.strip().startswith("|"):
                if current_type != "table":
                    flush()
                    current_type = "table"
                current.append(line.strip())
            elif not line.strip():
                flush()
                current_type = "paragraph"
            else:
                if current_type != "paragraph":
                    flush()
                    current_type = "paragraph"
                current.append(line)
        flush()
        if not blocks:
            raise PermanentIngestionError("Document contains no extractable text")
        return ParsedDocument("document", tuple(blocks))


class _SafeHtmlParser(HTMLParser):
    ignored = {"script", "style", "noscript", "iframe", "object", "embed"}
    block_tags = {
        "p": "paragraph",
        "li": "list",
        "blockquote": "quote",
        "pre": "code",
        "table": "table",
        **{f"h{i}": "heading" for i in range(1, 7)},
    }

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.depth = 0
        self.active: BlockType | None = None
        self.parts: list[str] = []
        self.items: list[tuple[BlockType, str, int | None]] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        del attrs
        tag = tag.lower()
        if tag in self.ignored:
            self.depth += 1
        if self.depth:
            return
        if tag in self.block_tags:
            self._flush()
            self.active = self.block_tags[tag]  # type: ignore[assignment]
            self.heading_level = int(tag[1]) if tag.startswith("h") else None
        elif tag in {"br", "tr"}:
            self.parts.append("\n")
        elif tag in {"td", "th"}:
            self.parts.append(" | ")

    def handle_endtag(self, tag: str) -> None:
        tag = tag.lower()
        if tag in self.ignored and self.depth:
            self.depth -= 1
            return
        if not self.depth and tag in self.block_tags:
            self._flush()

    def handle_data(self, data: str) -> None:
        if not self.depth and self.active is not None:
            self.parts.append(data)

    def close(self) -> None:
        super().close()
        self._flush()

    def _flush(self) -> None:
        text = " ".join("".join(self.parts).split())
        if self.active and text:
            self.items.append((self.active, text, getattr(self, "heading_level", None)))
        self.parts = []
        self.active = None


class HtmlDocumentParser(ParserAdapter):
    extensions = (".html", ".htm")
    content_types = ("text/html", "application/xhtml+xml")

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        text = _decode_utf8(data, "HTML")
        parser = _SafeHtmlParser()
        try:
            parser.feed(text)
            parser.close()
        except Exception as exc:
            raise PermanentIngestionError("HTML document could not be parsed") from exc
        headings: list[str] = []
        blocks: list[KnowledgeBlock] = []
        for block_type, value, level in parser.items:
            if block_type == "heading" and level:
                headings[level - 1 :] = [value]
            blocks.append(
                _block(
                    file_name,
                    len(blocks),
                    block_type,
                    value,
                    section_path=tuple(headings),
                    heading_context=headings[-1] if headings else None,
                )
            )
        if not blocks:
            raise PermanentIngestionError("Document contains no extractable text")
        return ParsedDocument("document", tuple(blocks))


def _column_name(value: Any, index: int, seen: dict[str, int]) -> str:
    base = re.sub(r"\s+", "_", str(value).strip()).strip("_").lower() if value is not None else ""
    base = base or f"column_{index + 1}"
    seen[base] = seen.get(base, 0) + 1
    return base if seen[base] == 1 else f"{base}_{seen[base]}"


def _infer_type(values: list[Any]) -> str:
    present = [value for value in values if value is not None and value != ""]
    if not present:
        return "string"
    checks = (
        ("boolean", lambda value: isinstance(value, bool)),
        ("integer", lambda value: isinstance(value, int) and not isinstance(value, bool)),
        (
            "decimal",
            lambda value: isinstance(value, (int, float, Decimal)) and not isinstance(value, bool),
        ),
        ("datetime", lambda value: isinstance(value, datetime)),
        ("date", lambda value: isinstance(value, date)),
    )
    for name, check in checks:
        if all(check(value) for value in present):
            return name
    return "string"


def _excel_column(index: int) -> str:
    result = ""
    while index:
        index, remainder = divmod(index - 1, 26)
        result = chr(65 + remainder) + result
    return result


def _tables_from_rows(
    file_name: str, sheet_name: str, rows: list[list[Any]]
) -> tuple[NormalizedTable, ...]:
    row_regions: list[tuple[int, list[list[Any]]]] = []
    current: list[list[Any]] = []
    start = 1
    for row_number, row in enumerate(rows, start=1):
        if any(value is not None and str(value).strip() for value in row):
            if not current:
                start = row_number
            current.append(row)
        elif current:
            row_regions.append((start, current))
            current = []
    if current:
        row_regions.append((start, current))

    regions: list[tuple[int, int, list[list[Any]]]] = []
    for start_row, row_region in row_regions:
        width = max(len(row) for row in row_region)
        active = [
            any(
                index < len(row) and row[index] is not None and str(row[index]).strip()
                for row in row_region
            )
            for index in range(width)
        ]
        band_start: int | None = None
        for index, has_value in enumerate([*active, False]):
            if has_value and band_start is None:
                band_start = index
            elif not has_value and band_start is not None:
                regions.append(
                    (start_row, band_start + 1, [row[band_start:index] for row in row_region])
                )
                band_start = None

    tables: list[NormalizedTable] = []
    for table_index, (start_row, start_column, region) in enumerate(regions, start=1):
        width = max(len(row) for row in region)
        if width > 2_000:
            raise PermanentIngestionError("Spreadsheet exceeds the column safety limit")
        padded = [row + [None] * (width - len(row)) for row in region]
        seen: dict[str, int] = {}
        names = [_column_name(value, index, seen) for index, value in enumerate(padded[0])]
        data_rows = padded[1:]
        columns = tuple(
            NormalizedColumn(name, _infer_type([row[index] for row in data_rows]))  # type: ignore[arg-type]
            for index, name in enumerate(names)
        )
        normalized_rows = tuple(
            NormalizedRow(
                row_number=start_row + offset,
                cell_range=(
                    f"{_excel_column(start_column)}{start_row + offset}:"
                    f"{_excel_column(start_column + width - 1)}{start_row + offset}"
                ),
                values={name: row[index] for index, name in enumerate(names)},
            )
            for offset, row in enumerate(data_rows, start=1)
            if any(value is not None and str(value).strip() for value in row)
        )
        if not normalized_rows:
            continue
        if len(normalized_rows) > 250_000:
            raise PermanentIngestionError("Spreadsheet exceeds the row safety limit")
        table_id = hashlib.sha256(
            f"{file_name}:{sheet_name}:{start_row}:{table_index}".encode()
        ).hexdigest()[:24]
        tables.append(
            NormalizedTable(
                table_id=table_id,
                sheet_name=sheet_name,
                cell_range=(
                    f"{_excel_column(start_column)}{start_row}:"
                    f"{_excel_column(start_column + width - 1)}{start_row + len(region) - 1}"
                ),
                columns=columns,
                rows=normalized_rows,
            )
        )
    return tuple(tables)


def _spreadsheet_document(file_name: str, tables: tuple[NormalizedTable, ...]) -> ParsedDocument:
    blocks: list[KnowledgeBlock] = []
    for table in tables:
        summary = f"Sheet: {table.sheet_name}; Range: {table.cell_range}; Columns: " + ", ".join(
            f"{column.name} ({column.inferred_type})" for column in table.columns
        )
        blocks.append(
            _block(
                file_name,
                len(blocks),
                "table",
                summary,
                sheet_name=table.sheet_name,
                cell_range=table.cell_range,
                table_id=table.table_id,
                metadata={"record_kind": "table_summary"},
            )
        )
        for row in table.rows:
            text = f"Sheet: {table.sheet_name}; Range: {row.cell_range}; " + "; ".join(
                f"{name}: {value}" for name, value in row.values.items() if value not in (None, "")
            )
            blocks.append(
                _block(
                    file_name,
                    len(blocks),
                    "row",
                    text,
                    sheet_name=table.sheet_name,
                    cell_range=row.cell_range,
                    table_id=table.table_id,
                    metadata={"record_kind": "row"},
                )
            )
    if not blocks:
        raise PermanentIngestionError("Spreadsheet contains no logical tables")
    return ParsedDocument("spreadsheet", tuple(blocks), tables)


class CsvParser(ParserAdapter):
    extensions = (".csv",)
    content_types = ("text/csv", "application/csv", "application/vnd.ms-excel")

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        text = _decode_utf8(data, "CSV")
        try:
            sample = "\n".join(text[:8192].splitlines()[:50])
            first_region = re.split(r"\n\s*\n", sample, maxsplit=1)[0]
            dialect = csv.Sniffer().sniff(first_region, delimiters=",;\t")
            rows = [
                [_parse_csv_value(value) for value in row]
                for row in csv.reader(StringIO(text), dialect)
            ]
        except csv.Error as exc:
            raise PermanentIngestionError("CSV document could not be parsed") from exc
        tables = _tables_from_rows(file_name, PurePosixPath(file_name).stem, rows)
        return _spreadsheet_document(file_name, tables)


class XlsxParser(ParserAdapter):
    extensions = (".xlsx",)
    content_types = ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",)

    def parse(self, data: bytes, *, file_name: str) -> ParsedDocument:
        try:
            from openpyxl import load_workbook

            workbook = load_workbook(BytesIO(data), read_only=True, data_only=False)
        except ImportError as exc:
            raise PermanentIngestionError("XLSX parser dependency is not installed") from exc
        except Exception as exc:
            raise PermanentIngestionError("XLSX workbook could not be parsed") from exc
        tables: list[NormalizedTable] = []
        try:
            for sheet in workbook.worksheets:
                if sheet.sheet_state != "visible":
                    continue
                rows: list[list[Any]] = []
                formulas: dict[tuple[int, int], str] = {}
                for cells in sheet.iter_rows():
                    row: list[Any] = []
                    for cell in cells:
                        if cell.data_type == "f":
                            formulas[(cell.row, cell.column)] = str(cell.value)
                            row.append(None)
                        else:
                            row.append(cell.value)
                    rows.append(row)
                sheet_tables = _tables_from_rows(file_name, sheet.title, rows)
                for table in sheet_tables:
                    start_column = _column_index(re.split(r"\d", table.cell_range, maxsplit=1)[0])
                    updated_rows: list[NormalizedRow] = []
                    for normalized_row in table.rows:
                        expressions = {
                            column.name: formulas[(normalized_row.row_number, column_index)]
                            for column_index, column in enumerate(table.columns, start=start_column)
                            if (normalized_row.row_number, column_index) in formulas
                        }
                        updated_rows.append(
                            replace(
                                normalized_row,
                                formula_columns=tuple(expressions),
                                formula_expressions=expressions,
                            )
                        )
                    tables.append(replace(table, rows=tuple(updated_rows)))
        finally:
            workbook.close()
        return _spreadsheet_document(file_name, tuple(tables))


class DocumentTextExtractor:
    """Parser registry/factory facade for all supported digital formats."""

    def __init__(self, parsers: tuple[ParserAdapter, ...] | None = None):
        self._parsers = parsers or (
            PdfParser(),
            DocxParser(),
            MarkdownParser(),
            HtmlDocumentParser(),
            TextParser(),
            XlsxParser(),
            CsvParser(),
        )

    def parse(self, data: bytes, *, content_type: str, file_name: str) -> ParsedDocument:
        normalized = content_type.lower().split(";", 1)[0].strip()
        suffix = PurePosixPath(file_name.lower()).suffix
        for parser in self._parsers:
            if suffix in parser.extensions:
                allowed = set(parser.content_types) | {
                    "application/octet-stream",
                    "application/zip",
                }
                if normalized not in allowed and not (
                    isinstance(parser, (TextParser, MarkdownParser, CsvParser))
                    and normalized == "text/plain"
                ):
                    raise PermanentIngestionError("File extension and content type do not match")
                return parser.parse(data, file_name=file_name)
        raise PermanentIngestionError(f"Unsupported document content type: {content_type}")

    def extract(self, data: bytes, *, content_type: str, file_name: str) -> list[ExtractedPage]:
        parsed = self.parse(data, content_type=content_type, file_name=file_name)
        grouped: dict[int, list[str]] = {}
        for block in parsed.blocks:
            page = block.page_number or 1
            grouped.setdefault(page, []).append(block.text)
        return [ExtractedPage(page, "\n\n".join(texts)) for page, texts in sorted(grouped.items())]


def _parse_csv_value(value: str) -> Any:
    cleaned = value.strip()
    if cleaned == "":
        return None
    if re.fullmatch(r"-?(?:0|[1-9]\d*)", cleaned) and not (
        cleaned.startswith("0") and len(cleaned) > 1
    ):
        try:
            return int(cleaned)
        except ValueError:
            pass
    if re.fullmatch(r"-?(?:0|[1-9]\d*)\.\d+", cleaned):
        try:
            return Decimal(cleaned)
        except ValueError:
            pass
    try:
        return date.fromisoformat(cleaned)
    except ValueError:
        return cleaned


def _column_index(letters: str) -> int:
    result = 0
    for letter in letters:
        result = result * 26 + ord(letter) - 64
    return result
